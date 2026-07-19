"""
Generate a tailored resume from the files in resume_input/.

Contact info, education, and each employer's company/location/dates/title are
extracted deterministically from resume_input/in_profile* and in_resume* —
there's no tailoring judgment involved in those fields, so a single Claude
Code CLI call (using your logged-in subscription, not a billed API key) is
only asked to produce the parts that actually require judgment: a tailored
summary, per-employer bullets selected from true source material, and a
tailored skills list. The result is written to resume_create/ as .docx +
.pdf; any resume already there is moved to resume_archive/ (.docx only)
first.

Requires the Claude Code CLI to be installed and logged in (`claude /login`).

Usage:
    python generate_resume.py
"""

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv

INK = RGBColor(0x1F, 0x2A, 0x44)
MUTED = RGBColor(0x44, 0x44, 0x44)
RULE_COLOR = "1F2A44"

ROOT = Path(__file__).resolve().parent
INPUT_DIR = ROOT / "resume_input"
CREATE_DIR = ROOT / "resume_create"
ARCHIVE_DIR = ROOT / "resume_archive"

MODEL = os.environ.get("CLAUDE_MODEL", "sonnet")
EFFORT = os.environ.get("CLAUDE_EFFORT", "medium")

# Distilled from resume_best_practices.md (see that file for sources/citations).
# Kept as a short, prompt-ready checklist so the model doesn't have to
# re-derive best practices, and so the per-run prompt stays compact.
BEST_PRACTICES = """1. Quantify impact wherever the source material has a real number — dollars, %, counts, time.
2. Start bullets with strong active verbs; never passive "was responsible for" phrasing.
3. Mirror the job posting's exact terminology for skills/title language the candidate honestly has.
4. Lead with the single most relevant, highest-impact bullet per role — recruiters scan, they don't read linearly at first.
5. Don't compress a 10+ year career to fit one page at the cost of cutting strong material; don't pad a short career to fill two.
6. Cut unquantified, generic duty-description bullets in favor of fewer, sharper, evidence-backed ones."""


def read_docx_text(path: Path) -> str:
    d = docx.Document(str(path))
    lines = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def find_latest(prefix: str) -> Path:
    matches = sorted(INPUT_DIR.glob(f"{prefix}*.docx"))
    if not matches:
        sys.exit(
            f"No file starting with '{prefix}' found in {INPUT_DIR}. "
            f"Expected a profile, prior resume, and job posting (in_profile*, "
            f"in_resume*, in_job*)."
        )
    return matches[-1]


# ---------------------------------------------------------------------------
# Deterministic extraction — no LLM involved, and none needed: contact info,
# education, and each job's company/location/dates/title carry no tailoring
# judgment, so parsing them straight out of the source .docx files is both
# faster and more accurate than asking a model to reproduce them.
# ---------------------------------------------------------------------------

PHONE_RE = re.compile(r"\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}")


def parse_applicant_info(profile_path: Path) -> dict:
    doc = docx.Document(str(profile_path))
    paras = [(p.style.name, p.text.strip()) for p in doc.paragraphs if p.text.strip()]

    lines = []
    in_section = False
    for style, text in paras:
        if not in_section:
            if text.lower().rstrip(":") in ("applicant info", "contact", "contact info"):
                in_section = True
            continue
        if style == "Normal" and text.lower().rstrip(":") not in (
            "applicant info",
            "contact",
            "contact info",
        ):
            break
        lines.append(text)

    email = next((line for line in lines if "@" in line), "")
    phone = next((line for line in lines if PHONE_RE.search(line)), "")
    remaining = [line for line in lines if line != email and line != phone]

    first_name, last_name, location = "", "", ""
    if remaining:
        parts = remaining[0].split()
        first_name = parts[0] if parts else ""
        last_name = " ".join(parts[1:]) if len(parts) > 1 else ""
    if len(remaining) > 1:
        location = ", ".join(remaining[1:])

    if not first_name or not email:
        sys.exit(
            f"Could not find name/email under an 'Applicant info' section in {profile_path.name}. "
            "Expected a section with the applicant's name, email, phone, and location as separate lines."
        )

    return {
        "first_name": first_name,
        "last_name": last_name,
        "email": email,
        "phone": phone,
        "location": location,
    }


def _split_header_line(text: str) -> tuple:
    """Split a 'Name | Location <tabs> Dates' line into (name, location, dates)."""
    name, _, rest = text.partition("|")
    location, _, dates = rest.partition("\t")
    return name.strip(), location.strip(), dates.strip(" \t")


def _parse_dated_entries(paras: list) -> list:
    """Parse (style, text) pairs into entries of {name, location, dates, detail, bullets}."""
    entries = []
    i = 0
    while i < len(paras):
        style, text = paras[i]
        if style != "Normal" or "|" not in text:
            i += 1
            continue
        name, location, dates = _split_header_line(text)
        i += 1
        detail = ""
        if i < len(paras) and paras[i][0] == "Normal":
            detail = paras[i][1].strip()
            i += 1
        bullets = []
        while i < len(paras) and paras[i][0] == "List Paragraph":
            bullets.append(paras[i][1].strip())
            i += 1
        entries.append(
            {"name": name, "location": location, "dates": dates, "detail": detail, "bullets": bullets}
        )
    return entries


def parse_prior_resume(resume_path: Path) -> dict:
    doc = docx.Document(str(resume_path))
    paras = [(p.style.name, p.text) for p in doc.paragraphs if p.text.strip()]

    section_names = {"summary", "education", "experience", "skills", "projects", "awards"}
    sections = {}
    current = None
    for style, text in paras:
        key = text.strip().rstrip(":").strip().lower()
        if style == "Normal" and key in section_names:
            current = key
            sections[current] = []
            continue
        if current:
            sections[current].append((style, text))

    education_entries = _parse_dated_entries(sections.get("education", []))
    education = [
        {
            "institution": e["name"],
            "location": e["location"],
            "dates": e["dates"],
            "degree": e["detail"],
        }
        for e in education_entries
    ]

    job_entries = _parse_dated_entries(sections.get("experience", []))
    jobs = [
        {
            "company": e["name"],
            "location": e["location"],
            "dates": e["dates"],
            "title": e["detail"],
            "source_bullets": e["bullets"],
        }
        for e in job_entries
    ]

    if not jobs:
        sys.exit(
            f"Could not parse any employers out of the EXPERIENCE section of {resume_path.name}. "
            "Expected 'Company | Location <tab> Dates' header lines followed by a title line and "
            "bulleted List Paragraph entries."
        )

    return {"education": education, "jobs": jobs}


# ---------------------------------------------------------------------------
# LLM call — a single pass that only produces the parts requiring judgment:
# the tailored summary, per-employer bullets, and a tailored skills list.
# ---------------------------------------------------------------------------

TAILOR_SYSTEM_PROMPT = """You are an expert resume strategist. Your job is to make a candidate's \
truthful, existing career history land as strongly as possible for ONE specific job \
posting, so a recruiter reading it decides to call the candidate for a screening or \
interview. Tailoring to the posting is the main thing you are being judged on — a \
resume that just restates the source documents in the original order is a failure, \
even if every fact in it is true.

Ground truth, non-negotiable:
- Every accomplishment, metric, and skill you use must already exist in the PROFILE \
and/or PRIOR RESUME documents. Never invent, exaggerate, or imply experience, tools, \
metrics, or scope the candidate doesn't have. If the posting wants something the \
candidate has never done, do not claim it and do not imply it through word choice \
(e.g. don't call real attribution-modeling work "incrementality testing" unless the \
source material actually describes incrementality testing).
- Do not stretch the truth or put the applicant "over their skis" — tailoring means \
selection, reordering, and honest reframing of real experience, never fabrication.
- Never state a total years-of-experience figure higher than what the PRIOR RESUME says.
- Never transplant a detail, scope claim, or outcome from one employer's true material \
onto a different employer's bullets, even if it sounds plausible for both. Never add a \
scope, audience, or downstream-impact claim (e.g. who a result "informed" or was \
"used by") beyond what the source material states for that specific accomplishment.
- Every bullet listed under an employer must describe work done at that employer only. \
Never mention a different employer, or that employer's outcome, inside another \
employer's bullets — not even as a "previously did X at Y" aside.
- Never name a specific tool, platform, methodology, or term of art from the job \
posting that the candidate hasn't actually used or done, even as a comparison or \
analogy ("similar to X," "analogous to driving Y adoption," "directly applicable to \
Z," "the same approach used for W"). Naming an unfamiliar posting-specific term next \
to real work implies familiarity with it — that is exactly the kind of "over their \
skis" claim this rule forbids, even when it's technically phrased as a comparison \
rather than a direct claim.

Tailoring — weight this heavily:
- Before writing anything, identify the 5-8 things this specific posting most cares \
about (its "You Will" / "You Have" priorities, its named responsibilities, its own \
vocabulary).
- For each employer, pull from every TRUE bullet available across PROFILE and PRIOR \
RESUME for that employer (the profile often has more detail/metrics than the old \
resume's bullets) and select the ones that most directly serve this posting's \
priorities. Prefer fewer, sharper, highly relevant bullets over reproducing the \
source's full list — 3-5 strong bullets per role, cutting or shrinking whatever \
doesn't support this posting, leading each role with its most relevant bullet.
- Reword bullets in the posting's own vocabulary wherever that vocabulary accurately \
describes what the candidate truly did.
- Write the summary as a direct, confident pitch for this exact role: open with the \
candidate's single strongest true qualification against the posting's top priority, \
and use the posting's own framing (its title, its named priorities) wherever that \
framing is honestly supported by the source material.
- Build the skills list from a genuine mix of categories, pulling only what's actually \
in PROFILE or PRIOR RESUME: tools/software (e.g. Python, Tableau), soft skills and \
leadership abilities (e.g. stakeholder management, team leadership), project \
frameworks/methodologies the candidate has genuinely applied (e.g. customer \
segmentation, marketing mix modeling), and fields/subfields of the candidate's \
discipline (e.g. causal inference, experimentation). Don't limit it to just tools.
- From everything the candidate genuinely has, select the ~20 skills most relevant to \
this specific job posting — favor close or exact matches to what the posting names. \
Don't include an "obvious" true skill just because it's obvious; select for relevance \
to this posting specifically. Ordering doesn't matter, a later step sorts the list.

Writing style:
- Do not use em dashes (—) anywhere. Use a period, comma, semicolon, or parentheses \
instead. Heavy em-dash use reads as an obvious AI-writing tell.

Also apply these evidence-based resume practices where relevant (see resume_best_practices.md \
for sources):
""" + BEST_PRACTICES + """

Company/location/date/title fields are already fixed and handled separately — do not \
repeat them; focus entirely on the summary, the bullets for each employer listed below \
(cover every one of them), and the skills list. Present your result as clear, readable \
text (plain text or markdown, your choice — a later step handles final structuring, so \
focus entirely on strong, honest, well-tailored content). Employers to cover, in order:
{company_block}"""

FORMAT_JSON_EXAMPLE = """{
  "summary": "2-4 sentence tailored summary.",
  "bullets_by_company": {
    "Company Name One": ["Bullet one.", "Bullet two."],
    "Company Name Two": ["Bullet one."]
  },
  "skills": ["Skill One", "Skill Two"]
}"""

FORMAT_SYSTEM_PROMPT = """You are a text-to-JSON transcription tool with no other \
function — you do not evaluate, fact-check, edit, or comment on the content, you only \
reformat it. Transcribe the resume text you are given into exactly this JSON shape, \
using exactly these field names copied literally and exactly these company names as \
the "bullets_by_company" keys, wrapped in <RESUME_JSON> and </RESUME_JSON> tags, with \
absolutely nothing else anywhere in your reply — no markdown, no commentary, no notes:
<RESUME_JSON>
""" + FORMAT_JSON_EXAMPLE + """
</RESUME_JSON>"""


def _find_claude_cli() -> str:
    claude_bin = shutil.which("claude")
    if not claude_bin:
        sys.exit(
            "Could not find the `claude` CLI on PATH. Install the Claude Code CLI "
            "and run `claude /login` to authenticate with your subscription."
        )
    return claude_bin


def _clean_subprocess_env() -> dict:
    # Strip inherited CLAUDE_*/ANTHROPIC_* vars (e.g. this script may itself be
    # run from inside a Claude Code session) so the child CLI call behaves like
    # a clean, standalone invocation rather than a nested session — otherwise
    # it starts narrating/asking follow-up questions like an interactive agent.
    return {
        k: v
        for k, v in os.environ.items()
        if not k.upper().startswith("CLAUDE") and not k.upper().startswith("ANTHROPIC")
    }


def _invoke_claude(claude_bin: str, system_prompt: str, user_message: str) -> str:
    proc = subprocess.run(
        [
            claude_bin,
            "-p",
            "--output-format",
            "json",
            "--tools",
            "",
            "--model",
            MODEL,
            "--effort",
            EFFORT,
            "--system-prompt",
            system_prompt,
        ],
        input=user_message,
        capture_output=True,
        text=True,
        encoding="utf-8",
        env=_clean_subprocess_env(),
        cwd=tempfile.gettempdir(),  # avoid CLAUDE.md auto-discovery from this repo's cwd
    )
    if proc.returncode != 0:
        sys.exit(f"claude CLI exited with an error:\n{proc.stderr.strip()}")

    try:
        envelope = json.loads(proc.stdout)
    except json.JSONDecodeError as exc:
        sys.exit(f"Could not parse claude CLI output as JSON ({exc}):\n{proc.stdout}")

    if envelope.get("is_error"):
        sys.exit(f"claude CLI reported an error: {envelope.get('result')}")

    return envelope.get("result", "")


def _strip_code_fence(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()[1:]
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    return text


def _flatten_skills(value) -> list:
    if isinstance(value, list):
        return [str(v) for v in value]
    if isinstance(value, dict):
        flat = []
        for group in value.values():
            flat.extend(_flatten_skills(group))
        return flat
    return []


def _try_parse_format_json(raw: str, companies: list) -> dict | None:
    match = re.search(r"<RESUME_JSON>(.*?)</RESUME_JSON>", raw, re.DOTALL)
    text = _strip_code_fence(match.group(1) if match else raw)
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        return None
    if not isinstance(parsed, dict):
        return None

    summary = ""
    for key in ("summary", "professional_summary", "resume_summary"):
        if parsed.get(key):
            summary = parsed[key]
            break

    raw_bullets = parsed.get("bullets_by_company") or {}
    # Match case/whitespace-insensitively since the model copies company
    # names back rather than using a fixed schema.
    normalized = {re.sub(r"\s+", " ", k.strip().lower()): v for k, v in raw_bullets.items()}
    bullets_by_company = {}
    for company in companies:
        key = re.sub(r"\s+", " ", company.strip().lower())
        bullets_by_company[company] = list(normalized.get(key) or [])

    skills = []
    for key in ("skills", "core_skills", "skill_list"):
        if parsed.get(key):
            skills = _flatten_skills(parsed[key])
            break

    return {"summary": summary, "bullets_by_company": bullets_by_company, "skills": skills}


SKILLS_SECTION_HEADING_RE = re.compile(
    r"^(core competencies|skills|skills\s*&?\s*tools|technical\s+(skills|toolkit))\s*:?\s*$",
    re.IGNORECASE,
)


def _clean_heading_line(line: str) -> str:
    return re.sub(r"^#{1,6}\s*|\*+", "", line).strip().rstrip(":")


def _split_skill_items(text: str) -> list:
    """Split on "|" (the model's category-item separator), then on "," —
    but never split a comma that's inside parentheses, since skill names
    sometimes list examples in parens (e.g. "Uplift Modeling (CausalML,
    PyLift)"), which a naive comma-split would break apart."""
    items = []
    for chunk in text.split("|"):
        depth = 0
        current = ""
        for ch in chunk:
            if ch == "(":
                depth += 1
            elif ch == ")":
                depth = max(0, depth - 1)
            if ch == "," and depth == 0:
                items.append(current)
                current = ""
            else:
                current += ch
        items.append(current)
    return [i.strip() for i in items if i.strip()]


CATEGORY_LINE_RE = re.compile(r"^-?\s*\*\*([^*]+):\*\*\s*(.+)$")


def extract_skills_from_draft(draft: str) -> list:
    """Deterministically pull every skill/tool out of the draft's
    competencies section — the JSON-transcription step doesn't reliably
    flatten this. Primarily detects the categorized layout ("**Category:**
    item | item", optionally bulleted) structurally, wherever it appears,
    since the model uses a different heading each run ("CORE COMPETENCIES",
    "CORE STRENGTHS", ...) rather than matching on heading text. Falls back
    to a heading-based search for a flat, uncategorized list."""
    lines = draft.splitlines()

    items = []
    for line in lines:
        match = CATEGORY_LINE_RE.match(line.strip())
        if match:
            items.extend(_split_skill_items(match.group(2)))
    if items:
        return items

    start = None
    for i, line in enumerate(lines):
        if SKILLS_SECTION_HEADING_RE.match(_clean_heading_line(line)):
            start = i + 1
            break
    if start is None:
        return []

    for line in lines[start:]:
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^#{1,6}\s+\S", stripped) or re.match(r"^-{3,}$", stripped):
            break  # next section heading or a horizontal-rule divider
        content = re.sub(r"^[-*]\s+", "", stripped)  # drop a bullet marker, if any
        items.extend(_split_skill_items(content))
    return items


def tailor_content(profile_text: str, resume_text: str, job_text: str, jobs: list) -> dict:
    claude_bin = _find_claude_cli()
    companies = [job["company"] for job in jobs]
    company_block = "\n".join(companies)
    draft_system_prompt = TAILOR_SYSTEM_PROMPT.replace("{company_block}", company_block)

    user_message = f"""PROFILE (source of truth for accomplishments, skills, and details):
{profile_text}

PRIOR RESUME (source of truth for which bullets belong to which employer):
{resume_text}

JOB POSTING (tailor emphasis, ordering, and phrasing to this):
{job_text}"""

    print("  Drafting tailored content...")
    draft = _invoke_claude(claude_bin, draft_system_prompt, user_message)

    print("  Converting to structured data...")
    parsed = None
    raw = ""
    for attempt in range(3):
        raw = _invoke_claude(claude_bin, FORMAT_SYSTEM_PROMPT, draft)
        parsed = _try_parse_format_json(raw, companies)
        total_bullets = sum(len(b) for b in (parsed or {}).get("bullets_by_company", {}).values())
        if parsed and parsed["summary"] and total_bullets:
            break
        print(f"  Attempt {attempt + 1} did not return usable structured data, retrying...")
    else:
        sys.exit(f"Could not get usable structured data after 3 attempts. Last raw response:\n{raw}")

    # The JSON-transcription step doesn't reliably flatten a categorized
    # skills section (e.g. "**Leadership:** Team Building | Change
    # Management") into the flat list we need — sometimes it drops whole
    # categories, sometimes it returns nothing at all. Parsing the draft's
    # skills section directly is far more reliable, so prefer it whenever it
    # finds anything.
    draft_skills = extract_skills_from_draft(draft)
    if draft_skills:
        parsed["skills"] = draft_skills

    return parsed


# ---------------------------------------------------------------------------
# Deterministic validation/correction guards — cheap, precise checks that
# don't depend on LLM self-compliance.
# ---------------------------------------------------------------------------

YEARS_RE = re.compile(r"(\d+)\+?\s*years?\s+of\s+experience", re.IGNORECASE)


def fix_years_of_experience(summary: str, resume_text: str) -> str:
    source_match = YEARS_RE.search(resume_text)
    generated_match = YEARS_RE.search(summary)
    if not source_match or not generated_match:
        return summary
    source_years = int(source_match.group(1))
    generated_years = int(generated_match.group(1))
    if generated_years > source_years + 2:
        print(
            f"  Correcting inflated experience claim: '{generated_years}+ years' -> "
            f"'{source_years}+ years' (source resume states {source_years} years)"
        )
        start, end = generated_match.span(1)
        summary = summary[:start] + str(source_years) + summary[end:]
    return summary


# Generic first words that shouldn't be treated as a company's identifying
# keyword (would cause false-positive matches against unrelated bullets).
_GENERIC_COMPANY_WORDS = {"the", "american", "national", "united", "global", "first"}


def _company_keywords(name: str) -> list:
    """Distinctive name(s) to search for — the full name, and its first word
    if that word is specific enough to be a reliable signal on its own (e.g.
    "Allstate" from "Allstate Insurance Company")."""
    keywords = [name]
    first_word = name.split()[0] if name.split() else ""
    if len(first_word) >= 4 and first_word.lower() not in _GENERIC_COMPANY_WORDS:
        keywords.append(first_word)
    return keywords


def strip_cross_employer_mentions(bullets_by_company: dict) -> dict:
    companies = list(bullets_by_company.keys())
    other_keywords = {
        company: [
            (other, kw) for other in companies if other != company for kw in _company_keywords(other)
        ]
        for company in companies
    }
    cleaned = {}
    for company, bullets in bullets_by_company.items():
        kept = []
        for bullet in bullets:
            hit = next(
                (other for other, kw in other_keywords[company] if kw.lower() in bullet.lower()),
                None,
            )
            if hit:
                print(
                    f"  Warning: dropped a bullet under {company} that mentioned "
                    f"{hit} (cross-employer content is not allowed): {bullet[:80]}..."
                )
                continue
            kept.append(bullet)
        cleaned[company] = kept
    return cleaned


NUMBER_RE = re.compile(r"\$?\d[\d,]*(?:\.\d+)?\s*[%KMB]?\+?")


def warn_unsupported_numbers(text: str, label: str, source_text: str) -> None:
    source_digits = {re.sub(r"[^\d]", "", n) for n in NUMBER_RE.findall(source_text)}
    for match in NUMBER_RE.findall(text):
        digits = re.sub(r"[^\d]", "", match)
        if len(digits) < 2:
            continue  # skip single-digit numbers (e.g. "3 direct reports") — too noisy to check
        if digits not in source_digits:
            print(f"  Warning: '{match}' in {label} was not found in the source documents — please verify.")


def warn_unverified_skills(skills: list, source_text: str) -> None:
    source_lower = source_text.lower()
    for skill in skills:
        core = re.sub(r"\s*\([^)]*\)", "", skill).strip().lower()
        if core and core not in source_lower:
            print(f"  Warning: skill '{skill}' wasn't found verbatim in the source documents — please verify.")


ANALOGY_RE = re.compile(
    r"\b(analogous to|directly applicable to|the same \w+ (?:used|applied) for)\b",
    re.IGNORECASE,
)


def warn_analogy_phrasing(bullets_by_company: dict) -> None:
    for company, bullets in bullets_by_company.items():
        for bullet in bullets:
            if ANALOGY_RE.search(bullet):
                print(
                    f"  Warning: a bullet under {company} uses comparison phrasing that "
                    f"may overclaim relevance to unfamiliar terms — please review: {bullet[:100]}..."
                )


def remove_em_dashes(text: str) -> str:
    """Backstop for the "no em dashes" style rule — prompt compliance isn't
    guaranteed, so replace any that slip through with a comma (safe for the
    fragment-style clauses resume bullets/summaries use)."""
    text = re.sub(r"\s*—\s*", ", ", text)
    return re.sub(r",\s*,", ",", text)


def _skill_key(skill: str) -> str:
    return re.sub(r"\s*\([^)]*\)", "", skill).strip().lower()


def parse_profile_skills(profile_path: Path) -> dict:
    """Deterministically pull the flat skill/tool lists out of the profile's
    'Software/Tools' and 'Skills' sections (List Paragraph entries under those
    headers) — used as a truthful pool to backfill anything the model omits.
    Kept separate: tool names are unambiguous and safe to always list in full;
    the broader skills phrases benefit more from job-specific filtering."""
    doc = docx.Document(str(profile_path))
    paras = [(p.style.name, p.text.strip()) for p in doc.paragraphs if p.text.strip()]

    tool_section_names = {"software/tools", "software / tools", "tools"}
    skill_section_names = {"skills"}
    tools, skills = [], []
    current = None
    for style, text in paras:
        key = text.rstrip(":").strip().lower()
        if style == "Normal":
            if key in tool_section_names:
                current = tools
            elif key in skill_section_names:
                current = skills
            else:
                current = None
            continue
        if current is not None and style == "List Paragraph":
            current.append(text)
    return {"tools": tools, "skills": skills}


def merge_and_sort_skills(tailored_skills: list, profile_skills: dict, job_text: str) -> list:
    """The model curates the actual skills list (mix of tools, soft skills,
    frameworks, and DS subfields, picked for relevance to this posting). This
    is just a narrow safety net: force-include a profile tool/skill only when
    the job posting itself explicitly names it too (e.g. "Python" shouldn't
    be missing when both the profile and the posting mention it) — it doesn't
    dump the whole profile in, which would fight the model's curation. Merge,
    dedup, and sort alphabetically."""
    job_lower = job_text.lower()
    must_include = [
        s
        for s in profile_skills["tools"] + profile_skills["skills"]
        if _skill_key(s) and _skill_key(s) in job_lower
    ]

    seen = {}
    for skill in tailored_skills + must_include:
        key = _skill_key(skill)
        if key and key not in seen:
            seen[key] = skill
    return sorted(seen.values(), key=str.lower)


# ---------------------------------------------------------------------------
# Assembly + docx/pdf output (unchanged from prior version).
# ---------------------------------------------------------------------------


def slugify_name(first: str, last: str) -> str:
    def clean(part: str) -> str:
        part = part.strip().lower()
        part = re.sub(r"[^a-z0-9]+", "-", part)
        return part.strip("-")

    return f"{clean(first)}-{clean(last)}"


def _set_bottom_border(paragraph, color: str = RULE_COLOR, size: int = 6) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _build_styles(d: docx.Document) -> dict:
    styles = d.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0

    name_style = styles.add_style("ResumeName", WD_STYLE_TYPE.PARAGRAPH)
    name_style.base_style = normal
    name_style.font.size = Pt(22)
    name_style.font.bold = True
    name_style.font.color.rgb = INK
    name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_style.paragraph_format.space_after = Pt(2)

    contact_style = styles.add_style("ResumeContact", WD_STYLE_TYPE.PARAGRAPH)
    contact_style.base_style = normal
    contact_style.font.size = Pt(9.5)
    contact_style.font.color.rgb = MUTED
    contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_style.paragraph_format.space_after = Pt(10)

    section_style = styles.add_style("ResumeSection", WD_STYLE_TYPE.PARAGRAPH)
    section_style.base_style = normal
    section_style.font.size = Pt(11.5)
    section_style.font.bold = True
    section_style.font.color.rgb = INK
    section_style.font.all_caps = True
    section_style.paragraph_format.space_before = Pt(12)
    section_style.paragraph_format.space_after = Pt(4)

    entry_title_style = styles.add_style("ResumeEntryTitle", WD_STYLE_TYPE.PARAGRAPH)
    entry_title_style.base_style = normal
    entry_title_style.font.italic = True
    entry_title_style.font.size = Pt(10)
    entry_title_style.font.color.rgb = MUTED
    entry_title_style.paragraph_format.space_after = Pt(3)

    bullet_style = styles["List Bullet"]
    bullet_style.font.name = "Calibri"
    bullet_style.font.size = Pt(10)
    bullet_style.paragraph_format.left_indent = Inches(0.2)
    bullet_style.paragraph_format.space_after = Pt(3)
    bullet_style.paragraph_format.line_spacing = 1.0

    return {"section": section_style, "entry_title": entry_title_style}


def _add_entry_header(d: docx.Document, left_text: str, right_text: str):
    p = d.add_paragraph()
    content_width = (
        d.sections[0].page_width - d.sections[0].left_margin - d.sections[0].right_margin
    )
    p.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
    left_run = p.add_run(left_text)
    left_run.bold = True
    if right_text:
        right_run = p.add_run(f"\t{right_text}")
        right_run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(1)
    return p


def build_docx(data: dict, out_path: Path) -> None:
    d = docx.Document()

    section = d.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = _build_styles(d)

    d.add_paragraph(f"{data['first_name']} {data['last_name']}", style="ResumeName")
    d.add_paragraph(
        " | ".join(
            part for part in (data["email"], data["phone"], data["location"]) if part
        ),
        style="ResumeContact",
    )

    def add_heading(text: str) -> None:
        p = d.add_paragraph(text, style=styles["section"])
        _set_bottom_border(p)

    add_heading("Summary")
    d.add_paragraph(data["summary"])

    if data.get("education"):
        add_heading("Education")
        for edu in data["education"]:
            header = " | ".join(
                part for part in (edu.get("institution"), edu.get("location")) if part
            )
            _add_entry_header(d, header, edu.get("dates", ""))
            degree_p = d.add_paragraph(edu.get("degree", ""))
            degree_p.paragraph_format.space_after = Pt(6)

    if data.get("experience"):
        add_heading("Experience")
        for job in data["experience"]:
            header = " | ".join(
                part for part in (job.get("company"), job.get("location")) if part
            )
            _add_entry_header(d, header, job.get("dates", ""))
            d.add_paragraph(job.get("title", ""), style=styles["entry_title"])
            bullets = job.get("bullets", [])
            for i, bullet in enumerate(bullets):
                p = d.add_paragraph(bullet, style="List Bullet")
                if i == len(bullets) - 1:
                    p.paragraph_format.space_after = Pt(8)

    if data.get("skills"):
        add_heading("Skills")
        d.add_paragraph(", ".join(data["skills"]))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(out_path))


def convert_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    from docx2pdf import convert

    convert(str(docx_path), str(pdf_path))


def archive_existing_outputs() -> None:
    ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
    CREATE_DIR.mkdir(parents=True, exist_ok=True)

    archive_time = datetime.now().strftime("%H-%M-%S")
    for docx_file in CREATE_DIR.glob("out_resume_*.docx"):
        archived_name = f"{docx_file.stem}-{archive_time}{docx_file.suffix}"
        shutil.move(str(docx_file), str(ARCHIVE_DIR / archived_name))
        print(f"Archived {docx_file.name} -> resume_archive/{archived_name}")

    for pdf_file in CREATE_DIR.glob("out_resume_*.pdf"):
        pdf_file.unlink()
        print(f"Removed superseded {pdf_file.name} (only .docx versions are archived)")


def main() -> None:
    load_dotenv(ROOT / ".env")

    profile_path = find_latest("in_profile")
    resume_path = find_latest("in_resume")
    job_path = find_latest("in_job")

    print(f"Profile: {profile_path.name}")
    print(f"Prior resume: {resume_path.name}")
    print(f"Job posting: {job_path.name}")

    contact = parse_applicant_info(profile_path)
    parsed_resume = parse_prior_resume(resume_path)

    profile_text = read_docx_text(profile_path)
    resume_text = read_docx_text(resume_path)
    job_text = read_docx_text(job_path)

    print(f"Tailoring content with {MODEL} (effort={EFFORT})...")
    tailored = tailor_content(profile_text, resume_text, job_text, parsed_resume["jobs"])

    summary = fix_years_of_experience(tailored["summary"], resume_text)
    summary = remove_em_dashes(summary)
    bullets_by_company = strip_cross_employer_mentions(tailored["bullets_by_company"])
    bullets_by_company = {c: [remove_em_dashes(b) for b in bs] for c, bs in bullets_by_company.items()}
    warn_analogy_phrasing(bullets_by_company)

    profile_skills = parse_profile_skills(profile_path)
    skills = merge_and_sort_skills(tailored["skills"], profile_skills, job_text)

    source_text = profile_text + "\n" + resume_text
    warn_unsupported_numbers(summary, "the summary", source_text)
    for company, bullets in bullets_by_company.items():
        for bullet in bullets:
            warn_unsupported_numbers(bullet, f"a bullet under {company}", source_text)
    warn_unverified_skills(skills, source_text)

    experience = [
        {
            "company": job["company"],
            "location": job["location"],
            "dates": job["dates"],
            "title": job["title"],
            "bullets": bullets_by_company.get(job["company"]) or job["source_bullets"],
        }
        for job in parsed_resume["jobs"]
    ]

    data = {
        **contact,
        "summary": summary,
        "education": parsed_resume["education"],
        "experience": experience,
        "skills": skills,
    }

    archive_existing_outputs()

    today = datetime.now().strftime("%Y-%m-%d")
    base_name = f"out_resume_{slugify_name(data['first_name'], data['last_name'])}_{today}"
    docx_path = CREATE_DIR / f"{base_name}.docx"
    pdf_path = CREATE_DIR / f"{base_name}.pdf"

    build_docx(data, docx_path)
    print(f"Wrote {docx_path.relative_to(ROOT)}")

    try:
        convert_to_pdf(docx_path, pdf_path)
        print(f"Wrote {pdf_path.relative_to(ROOT)}")
    except Exception as exc:  # docx2pdf requires MS Word via COM automation
        print(f"Warning: could not generate PDF ({exc}). The .docx was still created.")


if __name__ == "__main__":
    main()
